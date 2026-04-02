"""
renderers/docx.py
=================
Renders a SummaryTable to a Word .docx file (as bytes).

Table style follows common journal/policy paper conventions:
  - Serif font (Times New Roman 11pt)
  - Thick top/bottom borders, thin mid-border
  - Stat rows in 9pt
  - No vertical borders (academic convention)
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

from .._protocol import SummaryTable

if TYPE_CHECKING:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def render_docx(table: SummaryTable) -> bytes:
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for Word output.\n" "Install with: pip install python-docx"
        )
    doc = Document()

    # ── Style: default to Times New Roman throughout ──────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"  # type: ignore[attr-defined]
    style.font.size = Pt(11)  # type: ignore[attr-defined]

    # ── Title ─────────────────────────────────────────────────────────────
    if table.title:
        p = doc.add_paragraph(table.title)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

    # ── Table ─────────────────────────────────────────────────────────────
    n_data_cols = len(table.col_labels)
    total_cols = n_data_cols + 1  # label + data

    t = doc.add_table(rows=0, cols=total_cols)
    t.style = "Table Grid"

    # Remove all borders first; we'll add selectively
    _clear_table_borders(t)

    # Header row
    hdr = t.add_row()
    _set_top_border(hdr, thick=True)
    _set_bottom_border(hdr, thick=False)
    hdr.cells[0].text = ""
    for i, label in enumerate(table.col_labels, start=1):
        c = hdr.cells[i]
        c.text = label
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].runs[0].bold = True

    # Data rows
    separator_pending = False
    for row in table.rows:
        if row.is_separator:
            separator_pending = True
            continue

        tr = t.add_row()
        if separator_pending:
            _set_top_border(tr, thick=False)
            separator_pending = False

        # Label cell
        lbl_cell = tr.cells[0]
        lbl_cell.text = row.label
        if row.is_stat:
            _set_font(lbl_cell, size=Pt(9), color=RGBColor(0x44, 0x44, 0x44))

        # Data cells
        for i, val in enumerate(row.values, start=1):
            c = tr.cells[i]
            c.text = val
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row.is_stat:
                _set_font(c, size=Pt(9), color=RGBColor(0x44, 0x44, 0x44))

    # Bottom border on last row
    if t.rows:
        _set_bottom_border(t.rows[-1], thick=True)

    # ── Footer ────────────────────────────────────────────────────────────
    footer_parts = []
    if table.stars_legend:
        footer_parts.append(table.stars_legend)
    footer_parts.extend(table.notes)

    if footer_parts:
        p = doc.add_paragraph()
        for part in footer_parts:
            run = p.add_run(part + "  ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Serialise to bytes ────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Border helpers ────────────────────────────────────────────────────────────


def _clear_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_top_border(row, thick: bool = False) -> None:
    _set_row_border(row, "top", thick)


def _set_bottom_border(row, thick: bool = False) -> None:
    _set_row_border(row, "bottom", thick)


def _set_row_border(row, position: str, thick: bool) -> None:
    sz = "12" if thick else "6"  # half-points: 12 = 1.5pt, 6 = 0.75pt
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        border = OxmlElement(f"w:{position}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)


def _set_font(cell, size: Pt, color: RGBColor) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = size
            run.font.color.rgb = color
